import os
import sys

# Add scripts directory to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

os.chdir(os.path.dirname(os.path.abspath(__file__)))
os.environ['DOTENV_LOADED'] = '1'

from dotenv import load_dotenv
load_dotenv()

import requests
from io import BytesIO
from docx import Document
from docx.shared import Inch

url = 'https://maas-watermark-prod-new.cn-wlcb.ufileos.com/202604281739218b6fbe6966e444f3_watermark.png?UCloudPublicKey=TOKEN_6df395df-5d8c-4f69-90f8-a4fe46088958&Signature=GoNNO4kzH%2FO1XLMtvlIdEpiyIRM%3D&Expires=1777973970'

print(f"Testing image download from URL...")
print(f"URL length: {len(url)}")

try:
    response = requests.get(url, timeout=30, headers={'User-Agent': 'Mozilla/5.0'})
    print(f'Status: {response.status_code}')
    print(f'Content length: {len(response.content)} bytes')
    print(f'Content-Type: {response.headers.get("Content-Type", "unknown")}')

    # Create doc and add image
    doc = Document()
    doc.add_heading('Test Image Export', 0)
    doc.add_picture(BytesIO(response.content), width=6 * Inch)
    print("Image added to document")

    output_path = os.path.join(os.path.dirname(__file__), 'test_image_export.docx')
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

except Exception as e:
    print(f'Error: {type(e).__name__}: {e}')
    import traceback
    traceback.print_exc()