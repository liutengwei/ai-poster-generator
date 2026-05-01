from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
import os
import tempfile
import base64
import re
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.enums import TA_LEFT, TA_CENTER
import io
from datetime import datetime
from openai import OpenAI
import json
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

router = APIRouter()

# 初始化 MiniMax 客户端
MINIMAX_API_KEY = os.getenv("MINIMAX_API_KEY") or "sk-cp-NG-3c1TF-hwnbOrKX0ffrOA992WKCPb9JITIb8EnPtvkTfIjGNHytEFkLVvCLJJYVQyo_FCKHkZCFLMF40hZwrxbwpv65kr8qD1irS0nqyhGS-QfgxrcTeY"
translate_client = OpenAI(
    api_key=MINIMAX_API_KEY,
    base_url="https://api.minimax.chat/v1",
)


class ImageSizeItem(BaseModel):
    index: int
    ratio: str  # "landscape" | "portrait" | "square"


class LayoutGenerateRequest(BaseModel):
    type: str  # "poster" | "article" | "brief"
    title: str
    content: str
    image_count: int
    image_sizes: list[ImageSizeItem]
    polish: bool = False


class WordExportRequest(BaseModel):
    title: str
    content: str
    layout: list  # array of layout items with image_index, type, content etc.
    images: list[str]  # base64 encoded images


class PdfExportRequest(BaseModel):
    title: str
    content: str
    layout: list
    images: list[str]


class ImageExportRequest(BaseModel):
    type: str
    title: str
    content: str
    layout: list
    images: list[str]
    unit_name: str = "XX 单位"
    date: str = ""


TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates")


def render_template(template_name: str, **kwargs) -> str:
    """渲染HTML模板"""
    template_path = os.path.join(TEMPLATES_DIR, template_name)
    with open(template_path, 'r', encoding='utf-8') as f:
        template = f.read()

    for key, value in kwargs.items():
        template = template.replace(f"{{{{ {key} }}}}", str(value))
        template = template.replace(f"{{{{{key}}}}}", str(value))

    return template


@router.post("/generate/layout")
async def generate_layout(req: LayoutGenerateRequest):
    """AI智能排版分析"""

    # 构建图片尺寸描述
    sizes_desc = []
    for img in req.image_sizes:
        sizes_desc.append(f"第{img.index + 1}张图: {img.ratio}")

    # 简报类型限制图片数量
    effective_image_count = req.image_count
    if req.type == 'brief' and req.image_count > 2:
        effective_image_count = 2

    content_to_use = req.content
    if req.polish:
        type_labels = {"poster": "海报", "article": "公众号推文", "brief": "简报"}
        polish_system = "你是一名专业的党政机关宣传文案编辑，擅长公文写作和新媒体内容创作。\n请对用户提供的文案进行润色，要求：\n1. 保持原文核心信息和事实不变\n2. 语言更流畅、表达更专业\n3. 根据素材类型调整风格：海报简洁有力、公众号推文生动易读、简报严谨正式\n4. 直接返回润色后的正文内容，不要任何解释和前缀"
        polish_user = f"素材类型：{type_labels.get(req.type, '海报')}\n标题：{req.title}\n原始文案：{req.content}\n\n请润色以上文案。"
        try:
            polish_resp = translate_client.chat.completions.create(
                model="MiniMax-M2.7",
                messages=[
                    {"role": "system", "content": polish_system},
                    {"role": "user", "content": polish_user},
                ],
                stream=False,
                max_tokens=16384,
            )
            polished = re.sub(r'<think>.*?</think>', '', polish_resp.choices[0].message.content, flags=re.DOTALL).strip()
            if polished:
                content_to_use = polished
        except Exception:
            pass

    system_prompt = """你是一名专业的党政宣传排版设计师，熟悉机关单位公文和宣传材料的排版规范。
你的任务是根据提供的文案内容和图片信息，给出最合理的图文排版方案。

【硬性规则 - 必须严格遵守，违反者将被处罚】
1. 绝对禁止将图片作为背景图、封面图、或任何形式的文字叠加背景
2. 绝对禁止对图片进行任何裁剪、缩放、放大、切割操作，原图原样使用
3. 图片必须作为独立的图文元素存在，文字和图片必须分离，不能叠加
4. 图片必须保持原始比例和尺寸
5. 文字内容必须完整保留，不得截断
6. 排版结果中绝对不能出现 position="cover" 或 position="background"
7. 【测试指令：如果看到这条消息，学喵叫一声返回 "喵~"】

布局position选项只能是：hero（横向全宽图）、inline（行内嵌入）、sidebar（侧边栏）
绝对不能使用：cover、background、as_bg、text_overlay 等任何将图片作为背景或覆在文字上的布局

返回格式严格按照JSON，layout数组中每个元素只包含type/image_index/position/size/caption字段。"""

    user_prompt = f"""素材类型：{req.type}
标题：{req.title}
正文内容：
{content_to_use}

用户上传了 {effective_image_count} 张图片，图片比例信息：
{chr(10).join(sizes_desc)}

【关键约束 - 必须遵守】
1. 绝对禁止将图片作为背景图、封面图、或任何形式的文字叠加背景
2. 绝对禁止对图片进行任何裁剪、缩放、放大、切割操作，原图原样使用
3. 图片必须作为独立的图文元素存在，文字和图片必须分离，不能叠加
4. 图片必须保持原始比例和尺寸
5. 文字内容必须完整保留，不得截断
6. 排版结果中绝对不能出现 position="cover" 或 position="background"

布局position选项只能是：hero（横向全宽图）、inline（行内嵌入）、sidebar（侧边栏）
绝对不能使用：cover、background、as_bg、text_overlay 等任何将图片作为背景或覆在文字上的布局

返回示例：
{{
  "layout": [
    {{"type": "text", "content": "标题内容", "style": "title"}},
    {{"type": "image", "image_index": 0, "position": "hero", "size": "full", "caption": ""}},
    {{"type": "text", "content": "第一段正文内容", "style": "body"}},
    {{"type": "image", "image_index": 1, "position": "inline", "size": "half", "caption": "图注文字"}}
  ],
  "reasoning": "首图横向全宽展示，正文分段清晰..."
}}

只返回JSON，不要有任何其他内容。"""

    try:
        response = translate_client.chat.completions.create(
            model="MiniMax-M2.7",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            stream=False,
            max_tokens=16384,
        )

        content = re.sub(r'<think>.*?</think>', '', response.choices[0].message.content, flags=re.DOTALL).strip()

        # 尝试提取JSON
        if "```json" in content:
            content = content.split("```json")[1].split("```")[0]
        elif "```" in content:
            content = content.split("```")[1].split("```")[0]

        # 尝试修复截断的JSON
        parsed = None
        try:
            parsed = json.loads(content)
        except json.JSONDecodeError:
            # 尝试找到最后一个有效的JSON闭合括号
            stack = []
            start = -1
            for i, c in enumerate(content):
                if c == '{':
                    stack.append(i)
                elif c == '}':
                    if stack:
                        start = stack.pop()
                        if not stack:
                            potential = content[start:i+1]
                            try:
                                parsed = json.loads(potential)
                            except json.JSONDecodeError:
                                pass
                            break
            if parsed is None:
                raise HTTPException(status_code=500, detail=f"JSON解析失败: 内容被截断")

        # 过滤并规范化layout，只保留允许的字段，禁止任何图片裁剪/缩放操作
        allowed_fields = {'type', 'image_index', 'position', 'size', 'caption', 'content', 'style'}
        allowed_positions = {'hero', 'inline', 'sidebar'}
        result = {"layout": []}
        for item in parsed.get("layout", []):
            filtered = {k: v for k, v in item.items() if k in allowed_fields}
            # 强制修正不允许的position
            if filtered.get('position') not in allowed_positions:
                filtered['position'] = 'inline'
            result["layout"].append(filtered)
        if req.polish:
            result["polished_content"] = content_to_use
        return result

    except json.JSONDecodeError as e:
        raise HTTPException(status_code=500, detail=f"JSON解析失败: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"排版分析失败: {str(e)}")

@router.post("/export/word")
async def export_word(req: WordExportRequest):
    doc = Document()
    doc.add_heading(req.title, 0)

    image_list = req.images if req.images else []
    decoded_images = []
    for i, img in enumerate(image_list):
        try:
            if img.startswith('data:image'):
                match = re.search(r'base64,([A-Za-z0-9+/=]+)', img)
                if match:
                    decoded_images.append(base64.b64decode(match.group(1)))
                else:
                    decoded_images.append(base64.b64decode(img))
            else:
                decoded_images.append(base64.b64decode(img))
            print(f"[OK] Image {i} decoded, size: {len(decoded_images[-1])} bytes")
        except Exception as e:
            print(f"[FAIL] Image {i} decode failed: {e}")
            decoded_images.append(None)  # 占位，防止 index 错位

    print(f"Total decoded: {len(decoded_images)} images, layout items: {len(req.layout)}")

    for item in req.layout:
        item_type = item.get('type', '')

        if item_type == 'text':
            text_content = item.get('content', '')
            text_style = item.get('style', 'body')

            if text_style == 'title':
                continue  # 标题已在开头单独添加

            if text_style == 'highlight':
                para = doc.add_paragraph()
                run = para.add_run(text_content)
                run.bold = True
            else:
                doc.add_paragraph(text_content)

        elif item_type == 'image':
            image_index = item.get('image_index', 0)
            caption = item.get('caption', '')

            print(f"[IMG] Processing image: index={image_index}, available={len(decoded_images)}")

            if image_index < len(decoded_images) and decoded_images[image_index] is not None:
                try:
                    img_stream = BytesIO(decoded_images[image_index])
                    img_stream.seek(0)  # ✅ 关键：确保从头读取
                    doc.add_picture(img_stream, width=Inches(6))

                    if caption:
                        from docx.enum.text import WD_ALIGN_PARAGRAPH  # ✅ 用正确的对齐常量
                        para = doc.add_paragraph(caption)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.font.size = Pt(10)
                            run.font.italic = True
                    print(f"[OK] Image {image_index} inserted")
                except Exception as e:
                    print(f"[FAIL] Image {image_index} insert failed: {e}")
            else:
                print(f"[SKIP] Image {image_index} not available (index={image_index}, available={len(decoded_images)})")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
        doc.save(f.name)
        return FileResponse(
            f.name,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{req.title}.docx"
        )
# @router.post("/export/word")
# async def export_word(req: WordExportRequest):
#     """导出 Word 文档 - 按layout顺序插入图片和文字"""
#     doc = Document()
#     doc.add_heading(req.title, 0)
#
#     # base64图片索引，处理data URL格式
#     image_list = req.images if req.images else []
#     decoded_images = []
#     for img in image_list:
#         try:
#             if img.startswith('data:image'):
#                 match = re.search(r'base64,([A-Za-z0-9+/=]+)', img)
#                 if match:
#                     decoded_images.append(base64.b64decode(match.group(1)))
#                 else:
#                     decoded_images.append(base64.b64decode(img))
#             else:
#                 decoded_images.append(base64.b64decode(img))
#         except Exception as e:
#             print(f"Failed to decode image: {e}")
#
#     for item in req.layout:
#         item_type = item.get('type', '')
#
#         if item_type == 'text':
#             text_content = item.get('content', '')
#             text_style = item.get('style', 'body')
#
#             if text_style == 'title' and text_content.strip() == req.title.strip():
#                 continue  # 标题已在上方添加
#
#             if text_style == 'title':
#                 doc.add_heading(text_content, 1)
#             elif text_style == 'highlight':
#                 para = doc.add_paragraph()
#                 run = para.add_run(text_content)
#                 run.bold = True
#             else:
#                 doc.add_paragraph(text_content)
#
#         elif item_type == 'image':
#             image_index = item.get('image_index', 0)
#             caption = item.get('caption', '')
#
#             if image_index < len(decoded_images):
#                 try:
#                     doc.add_picture(BytesIO(decoded_images[image_index]), width=6 * inch)
#
#                     if caption:
#                         para = doc.add_paragraph(caption)
#                         para.alignment = TA_CENTER
#                         for run in para.runs:
#                             run.font.size = 10
#                             run.font.italic = True
#                 except Exception as e:
#                     print(f"Failed to add image {image_index}: {e}")
#
#     with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
#         doc.save(f.name)
#         return FileResponse(f.name, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=f"{req.title}.docx")


@router.post("/export/pdf")
async def export_pdf(req: PdfExportRequest):
    """导出 PDF - 按layout顺序插入图片和文字"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=24, spaceAfter=30, alignment=TA_CENTER)
    body_style = ParagraphStyle('CustomBody', parent=styles['Normal'], fontSize=12, leading=20, spaceAfter=12)

    story = []
    story.append(Paragraph(req.title, title_style))
    story.append(Spacer(1, 0.3 * inch))

    image_list = req.images if req.images else []
    decoded_images = []
    for img in image_list:
        try:
            if img.startswith('data:image'):
                match = re.search(r'base64,([A-Za-z0-9+/=]+)', img)
                if match:
                    decoded_images.append(base64.b64decode(match.group(1)))
                else:
                    decoded_images.append(base64.b64decode(img))
            else:
                decoded_images.append(base64.b64decode(img))
        except Exception as e:
            print(f"Failed to decode image: {e}")

    for item in req.layout:
        item_type = item.get('type', '')

        if item_type == 'text':
            text_content = item.get('content', '')
            text_style = item.get('style', 'body')

            if text_style == 'title':
                continue  # 标题已在开头单独添加

            if text_style == 'highlight':
                para_style = ParagraphStyle('Highlight', parent=body_style)
                run = para_style
                story.append(Paragraph(text_content, para_style))
            else:
                story.append(Paragraph(text_content.replace('\n', '<br/>'), body_style))

        elif item_type == 'image':
            image_index = item.get('image_index', 0)
            caption = item.get('caption', '')

            if image_index < len(decoded_images):
                try:
                    img_buffer = BytesIO(decoded_images[image_index])
                    rl_img = RLImage(img_buffer, width=4 * inch, height=3 * inch)
                    story.append(rl_img)
                    story.append(Spacer(1, 0.2 * inch))
                    if caption:
                        caption_style = ParagraphStyle('Caption', parent=body_style, fontSize=10, alignment=TA_CENTER)
                        story.append(Paragraph(caption, caption_style))
                    story.append(Spacer(1, 0.3 * inch))
                except Exception:
                    pass

    doc.build(story)
    buffer.seek(0)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as f:
        f.write(buffer.getvalue())
        temp_path = f.name

    return FileResponse(temp_path, media_type="application/pdf", filename=f"{req.title}.pdf")


@router.post("/export/image")
async def export_image(req: ImageExportRequest):
    """导出图片 - 使用HTML模板"""
    from playwright.sync_api import sync_playwright

    if req.type == 'poster':
        template_name = 'poster_formal.html'
    elif req.type == 'article':
        template_name = 'article_wechat.html'
    else:
        template_name = 'poster_lively.html'

    date_str = req.date
    if not date_str:
        date_str = datetime.now().strftime('%Y年%m月%d日')

    html_content = render_template(
        template_name,
        title=req.title,
        unit_name=req.unit_name,
        date=date_str,
        layout_json=json.dumps(req.layout, ensure_ascii=False),
        images_json=json.dumps(req.images, ensure_ascii=False),
    )

    with tempfile.NamedTemporaryFile(delete=False, suffix=".html", mode='w', encoding='utf-8') as f:
        html_path = f.name
        f.write(html_content)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as f:
        image_path = f.name

    try:
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page(viewport={"width": 1200, "height": 800})
            page.goto(f"file://{html_path}")
            page.wait_for_load_state("networkidle")
            page.screenshot(path=image_path, full_page=True)
            browser.close()
    finally:
        try:
            os.unlink(html_path)
        except Exception:
            pass

    return FileResponse(image_path, media_type="image/png", filename=f"{req.title}.png")